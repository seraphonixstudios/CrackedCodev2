"""Knowledge Base / RAG Upload v2.9.4 - Upload documents to augment RAG.

Upload PDFs, markdown files, text files, and code documentation to extend
the semantic search capabilities beyond just the codebase.

Usage:
    from src.knowledge_base import get_knowledge_base
    kb = get_knowledge_base()
    kb.upload_document("design_spec.pdf")
    results = kb.search("authentication flow")
"""

import json
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

from src.logger_config import get_logger

logger = get_logger("KnowledgeBase")


# ── Data Models ────────────────────────────────────────────────────────────

@dataclass
class Document:
    """A document in the knowledge base."""
    id: str
    title: str
    path: str
    content_type: str  # pdf, markdown, text, code, html
    chunks: List[str] = field(default_factory=list)
    embeddings: List[List[float]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    uploaded_at: str = ""
    size: int = 0


@dataclass
class SearchResult:
    """Search result from the knowledge base."""
    document_id: str
    title: str
    chunk_index: int
    content: str
    score: float
    metadata: Dict[str, Any] = field(default_factory=dict)


# ── Knowledge Base ─────────────────────────────────────────────────────────

class KnowledgeBase:
    """Document upload and semantic search for RAG augmentation."""
    
    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".md": "markdown",
        ".markdown": "markdown",
        ".txt": "text",
        ".rst": "text",
        ".html": "html",
        ".htm": "html",
        ".docx": "docx",
    }
    
    def __init__(self, storage_dir: str = ".crackedcode/knowledge"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        self.documents: Dict[str, Document] = {}
        self._load_index()
    
    def _load_index(self):
        """Load document index from disk."""
        index_path = self.storage_dir / "index.json"
        if index_path.exists():
            try:
                with open(index_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                for doc_data in data.get("documents", []):
                    doc = Document(
                        id=doc_data["id"],
                        title=doc_data["title"],
                        path=doc_data["path"],
                        content_type=doc_data["content_type"],
                        metadata=doc_data.get("metadata", {}),
                        uploaded_at=doc_data.get("uploaded_at", ""),
                        size=doc_data.get("size", 0),
                    )
                    self.documents[doc.id] = doc
                logger.info(f"Loaded {len(self.documents)} documents from knowledge base")
            except Exception as e:
                logger.warning(f"Failed to load knowledge base index: {e}")
    
    def _save_index(self):
        """Save document index to disk."""
        index_path = self.storage_dir / "index.json"
        data = {
            "documents": [
                {
                    "id": doc.id,
                    "title": doc.title,
                    "path": doc.path,
                    "content_type": doc.content_type,
                    "metadata": doc.metadata,
                    "uploaded_at": doc.uploaded_at,
                    "size": doc.size,
                }
                for doc in self.documents.values()
            ]
        }
        with open(index_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
    
    def upload_document(self, file_path: str, title: Optional[str] = None,
                        metadata: Optional[Dict[str, Any]] = None) -> Document:
        """Upload and index a document."""
        from datetime import datetime
        import hashlib
        
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Determine content type
        ext = path.suffix.lower()
        content_type = self.SUPPORTED_TYPES.get(ext, "text")
        
        # Generate document ID
        doc_id = hashlib.md5(str(path).encode()).hexdigest()[:12]
        
        # Read content
        content = self._extract_content(path, content_type)
        
        # Chunk content
        chunks = self._chunk_content(content)
        
        # Generate embeddings
        embeddings = self._generate_embeddings(chunks)
        
        # Create document
        doc = Document(
            id=doc_id,
            title=title or path.stem,
            path=str(path),
            content_type=content_type,
            chunks=chunks,
            embeddings=embeddings,
            metadata=metadata or {},
            uploaded_at=datetime.utcnow().isoformat(),
            size=len(content),
        )
        
        self.documents[doc_id] = doc
        self._save_index()
        
        # Save chunks
        self._save_document_chunks(doc)
        
        logger.info(f"Uploaded document: {doc.title} ({len(chunks)} chunks)")
        return doc
    
    def _extract_content(self, path: Path, content_type: str) -> str:
        """Extract text content from a file."""
        if content_type == "pdf":
            return self._extract_pdf(path)
        elif content_type == "docx":
            return self._extract_docx(path)
        elif content_type == "html":
            return self._extract_html(path)
        else:
            # Text-based files
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    
    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF."""
        try:
            import PyPDF2
            text = ""
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
            return text
        except ImportError:
            logger.warning("PyPDF2 not installed, cannot extract PDF")
            return ""
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""
    
    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX."""
        try:
            import docx
            doc = docx.Document(path)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            logger.warning("python-docx not installed, cannot extract DOCX")
            return ""
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""
    
    def _extract_html(self, path: Path) -> str:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            with open(path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                return soup.get_text(separator="\n")
        except ImportError:
            logger.warning("BeautifulSoup not installed, cannot extract HTML")
            return ""
        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            return ""
    
    def _chunk_content(self, content: str, chunk_size: int = 500,
                       overlap: int = 50) -> List[str]:
        """Split content into semantic chunks."""
        if not content:
            return []
        
        chunks = []
        # Split by paragraphs first
        paragraphs = content.split("\n\n")
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            if len(current_chunk) + len(para) < chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # If no chunks (e.g., very short content), use sliding window
        if not chunks and content:
            for i in range(0, len(content), chunk_size - overlap):
                chunk = content[i:i + chunk_size]
                if chunk.strip():
                    chunks.append(chunk.strip())
        
        return chunks
    
    def _generate_embeddings(self, chunks: List[str]) -> List[List[float]]:
        """Generate embeddings for chunks."""
        embeddings = []
        
        for chunk in chunks:
            try:
                embedding = self._get_embedding(chunk)
                embeddings.append(embedding)
            except Exception as e:
                logger.warning(f"Embedding generation failed: {e}")
                embeddings.append([0.0] * 768)  # Default zero embedding
        
        return embeddings
    
    def _get_embedding(self, text: str) -> List[float]:
        """Get embedding for text using Ollama or fallback."""
        try:
            from src.codebase_rag import EmbeddingProvider
            provider = EmbeddingProvider()
            return provider.get_embedding(text)
        except Exception:
            # Simple fallback: TF-IDF-like sparse embedding
            return self._fallback_embedding(text)
    
    def _fallback_embedding(self, text: str, dim: int = 768) -> List[float]:
        """Generate a simple hash-based fallback embedding."""
        import hashlib
        
        # Use hash to create deterministic vector
        vec = [0.0] * dim
        text_bytes = text.encode("utf-8")
        
        for i in range(0, min(len(text_bytes), dim * 4), 4):
            idx = (i // 4) % dim
            val = int.from_bytes(text_bytes[i:i+4], "little") / (2**31)
            vec[idx] += val
        
        # Normalize
        norm = sum(x**2 for x in vec) ** 0.5
        if norm > 0:
            vec = [x / norm for x in vec]
        
        return vec
    
    def _save_document_chunks(self, doc: Document):
        """Save document chunks to disk."""
        doc_dir = self.storage_dir / doc.id
        doc_dir.mkdir(exist_ok=True)
        
        # Save chunks
        chunks_path = doc_dir / "chunks.json"
        with open(chunks_path, "w", encoding="utf-8") as f:
            json.dump({"chunks": doc.chunks}, f)
        
        # Save embeddings
        embeddings_path = doc_dir / "embeddings.json"
        with open(embeddings_path, "w", encoding="utf-8") as f:
            json.dump({"embeddings": doc.embeddings}, f)
    
    def search(self, query: str, top_k: int = 5,
               filters: Optional[Dict[str, Any]] = None) -> List[SearchResult]:
        """Search the knowledge base for relevant documents."""
        if not self.documents:
            return []
        
        # Generate query embedding
        query_embedding = self._get_embedding(query)
        
        # Score all chunks
        results = []
        for doc in self.documents.values():
            # Apply filters
            if filters:
                skip = False
                for key, val in filters.items():
                    if doc.metadata.get(key) != val:
                        skip = True
                        break
                if skip:
                    continue
            
            for i, (chunk, embedding) in enumerate(zip(doc.chunks, doc.embeddings)):
                score = self._cosine_similarity(query_embedding, embedding)
                results.append(SearchResult(
                    document_id=doc.id,
                    title=doc.title,
                    chunk_index=i,
                    content=chunk[:300],
                    score=score,
                    metadata=doc.metadata,
                ))
        
        # Sort by score and return top_k
        results.sort(key=lambda r: r.score, reverse=True)
        return results[:top_k]
    
    def _cosine_similarity(self, a: List[float], b: List[float]) -> float:
        """Calculate cosine similarity between two vectors."""
        import math
        
        dot = sum(x * y for x, y in zip(a, b))
        norm_a = math.sqrt(sum(x**2 for x in a))
        norm_b = math.sqrt(sum(x**2 for x in b))
        
        if norm_a == 0 or norm_b == 0:
            return 0.0
        
        return dot / (norm_a * norm_b)
    
    def delete_document(self, doc_id: str) -> bool:
        """Delete a document from the knowledge base."""
        if doc_id not in self.documents:
            return False
        
        del self.documents[doc_id]
        self._save_index()
        
        # Remove chunks
        doc_dir = self.storage_dir / doc_id
        if doc_dir.exists():
            import shutil
            shutil.rmtree(doc_dir)
        
        logger.info(f"Deleted document: {doc_id}")
        return True
    
    def list_documents(self) -> List[Document]:
        """List all documents in the knowledge base."""
        return list(self.documents.values())
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get a document by ID."""
        return self.documents.get(doc_id)
    
    def get_stats(self) -> Dict[str, Any]:
        """Get knowledge base statistics."""
        total_docs = len(self.documents)
        total_chunks = sum(len(d.chunks) for d in self.documents.values())
        total_size = sum(d.size for d in self.documents.values())
        
        type_counts = {}
        for doc in self.documents.values():
            type_counts[doc.content_type] = type_counts.get(doc.content_type, 0) + 1
        
        return {
            "total_documents": total_docs,
            "total_chunks": total_chunks,
            "total_size_bytes": total_size,
            "documents_by_type": type_counts,
            "storage_dir": str(self.storage_dir),
        }


def get_knowledge_base(storage_dir: str = ".crackedcode/knowledge") -> KnowledgeBase:
    """Get the global knowledge base."""
    return KnowledgeBase(storage_dir=storage_dir)
